#!/usr/bin/env python3
# coding: utf-8
"""
validate.py -- nickel-ore skill report validator
Usage:
  python validate.py <report.docx>      validate a generated report
  python validate.py --check-skill      validate skill directory structure
"""
import sys
from pathlib import Path

SKILL_NAME = "nickel-ore"
SKILL_ROOT = Path(__file__).parent.parent

REQUIRED_SECTIONS = [
    ("项目概述", "project overview"),
    ("地质背景", "geological background"),
    ("矿床类型", "deposit type"),
    ("冶炼工艺", "processing"),
    ("资源潜力", "resource potential"),
    ("风险", "risk"),
]

REQUIRED_SKILL_FILES = [
    "SKILL.md",
    "knowledge/deposit_types.md",
    "knowledge/grade_tonnage.md",
    "knowledge/global_belts.md",
    "knowledge/processing_routes.md",
    "workflows/write_report.md",
    "workflows/quality_review.md",
    "workflows/analyze_project.md",
]

MIN_PARAGRAPHS = 50
MIN_SIZE_KB = 80
PLACEHOLDER_PATTERNS = ["[XXX]", "[待填写]", "[TBD]", "PLACEHOLDER", "（待补充）"]


def check_skill_structure():
    issues = []
    for rel in REQUIRED_SKILL_FILES:
        if not (SKILL_ROOT / rel).exists():
            issues.append(f"MISSING: {rel}")
    skill_md = SKILL_ROOT / "SKILL.md"
    if skill_md.exists():
        txt = skill_md.read_text(encoding="utf-8")
        # G24: frontmatter compressed, old trigger sections removed
        if "背景与定位" not in txt:
            issues.append("WARN: SKILL.md missing '背景与定位' (Thesis) section")
        if "Skill 误用反模式" not in txt:
            issues.append("WARN: SKILL.md missing 'Skill 误用反模式' section")
    return issues


def validate_report(report_path):
    issues, warnings = [], []
    path = Path(report_path)

    if not path.exists():
        print(f"FAIL: file not found -> {report_path}")
        sys.exit(1)

    size_kb = path.stat().st_size / 1024
    if size_kb < MIN_SIZE_KB:
        issues.append(f"file too small ({size_kb:.0f} KB < {MIN_SIZE_KB} KB)")

    try:
        from docx import Document
    except ImportError:
        print("FAIL: python-docx not installed -> pip install python-docx")
        sys.exit(1)

    try:
        doc = Document(str(path))
    except Exception as e:
        print(f"FAIL: cannot open DOCX -> {e}")
        sys.exit(1)

    paras = [p for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(p.text for p in doc.paragraphs)
    lower_text = full_text.lower()

    if len(paras) < MIN_PARAGRAPHS:
        issues.append(f"too few paragraphs ({len(paras)} < {MIN_PARAGRAPHS})")

    for cn, en in REQUIRED_SECTIONS:
        if cn not in full_text and en not in lower_text:
            issues.append(f"missing required section: {cn} / {en}")

    # nickel-ore specific: laterite vs magmatic type distinction + process route
    ni_type_ok = any(kw in full_text for kw in [
        "红土", "岩浆", "laterite", "magmatic", "HPAL", "RKEF", "Ferronickel",
        "硫化物", "sulfide", "komatiite", "Ni/Fe", "Class 1", "Class 2"
    ])
    if not ni_type_ok:
        issues.append("no nickel deposit type or process route indicators found (红土/岩浆/HPAL/RKEF)")

    fffd = full_text.count("\ufffd")
    if fffd > 0:
        issues.append(f"encoding error: {fffd} U+FFFD replacement chars found")

    hits = [p for p in PLACEHOLDER_PATTERNS if p in full_text]
    if hits:
        issues.append(f"unreplaced placeholders: {hits}")

    fig_count = len(doc.inline_shapes)
    if fig_count == 0:
        warnings.append("no embedded figures detected")
    elif size_kb < 1000 and fig_count > 3:
        warnings.append(f"{fig_count} figures but file is small ({size_kb:.0f} KB)")

    risk_ok = any(kw in full_text for kw in ["风险分析", "主要风险", "不确定", "risk analysis", "key risks"])
    if not risk_ok:
        issues.append("no risk analysis section found (required by anti-bias rule)")

    print(f"\n{'='*56}")
    print(f"  {SKILL_NAME} report validator")
    print(f"  file : {path.name}  ({size_kb:.0f} KB)")
    print(f"  paras: {len(paras)}   figures: {fig_count}")
    print(f"{'='*56}")
    if issues:
        print(f"  FAIL -- {len(issues)} issue(s):")
        for i, iss in enumerate(issues, 1):
            print(f"    {i}. {iss}")
    else:
        print("  PASS -- all checks passed")
    if warnings:
        print(f"  WARN ({len(warnings)}):")
        for w in warnings:
            print(f"    * {w}")
    print(f"{'='*56}\n")
    sys.exit(0 if not issues else 1)


def main():
    if len(sys.argv) < 2:
        print("usage: python validate.py <report.docx>")
        print("       python validate.py --check-skill")
        sys.exit(1)
    if sys.argv[1] == "--check-skill":
        issues = check_skill_structure()
        if issues:
            print(f"{SKILL_NAME} structure check -- FAIL:")
            for iss in issues:
                print(f"  * {iss}")
            sys.exit(1)
        else:
            print(f"{SKILL_NAME} structure check -- PASS")
    else:
        validate_report(sys.argv[1])


if __name__ == "__main__":
    main()
